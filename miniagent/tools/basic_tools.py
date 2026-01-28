import os
import re
import sys
import json
import math
import logging
import datetime
import platform
import subprocess
import webbrowser
from pathlib import Path
from typing import Dict, List, Any, Optional

import psutil
import requests

sys.path.append(os.getcwd())

from miniagent.tools import register_tool

logger = logging.getLogger(__name__)

@register_tool
def env_get(name: str) -> str:
    """
    Get an environment variable value.
    
    Args:
        name: Name of the environment variable
        
    Returns:
        Value of the environment variable or empty string if not found
    """
    return os.environ.get(name, "")


@register_tool
def env_set(name: str, value: str) -> str:
    """
    Set an environment variable (for current process only).
    
    Args:
        name: Name of the environment variable
        value: Value to set
        
    Returns:
        Status message
    """
    os.environ[name] = value
    return f"Set environment variable: {name}={value}"

@register_tool
def calculator(expression: str) -> float:
    """
     Calculate the result of a mathematical expression.

     Args:
         expression (str): The mathematical expression to evaluate.

     Returns:
         float: The result of the expression.
    """
    logger.info(f"[tool calls] calculator expression: {expression}")
    # Replace mathematical function names with methods from the math module
    math_names = {
        'sin': math.sin,
        'cos': math.cos,
        'tan': math.tan,
        'sqrt': math.sqrt,
        'pow': math.pow,
        'exp': math.exp,
        'log': math.log,
        'log10': math.log10,
        'pi': math.pi,
        'e': math.e
    }

    # For safety, clean unsafe characters from the expression
    expression = re.sub(r'[^\d+\-*/().a-zA-Z]', '', expression.strip())

        # Execute calculation
    try:
        # Use eval to calculate the expression, providing a safe context
        result = eval(expression, {"__builtins__": {}}, math_names)
        return float(result)
    except Exception as e:
        raise ValueError(f"Failed to calculate expression '{expression}': {str(e)}")

@register_tool
def get_current_time() -> Dict[str, Any]:
    """
    Get the current time in a dictionary format.

    Returns:
        Dict[str, Any]: A dictionary containing the current time with keys 'year', 'month', 'day', 'hour', 'minute', 'second'.
    """
    logger.info(f"[tool calls] get_current_time")
    current_time = datetime.datetime.now()
    return {
        'iso': current_time.isoformat(),
        'year': current_time.year,
        'month': current_time.month,
        'day': current_time.day,
        'hour': current_time.hour,
        'minute': current_time.minute,
        'second': current_time.second,
        "weekday": current_time.strftime("%A"),
        "formatted": current_time.strftime("%Y-%m-%d %H:%M:%S")
    }

@register_tool
def get_system_info() -> Dict[str, Any]:
    """
    Get detailed information about the system.
    
    Returns:
        Dictionary with system information like OS, version, architecture, etc.
    """
    logger.info(f"[tool calls] get_system_info")
    try:
        sys_info =  {
            "os": platform.system(),
            "os_release": platform.release(),
            "os_version": platform.version(),
            "architecture": platform.machine(),
            "processor": platform.processor(),
            "hostname": platform.node(),
            "python_version": platform.python_version(),
            "time": datetime.datetime.now().isoformat()
        }

        if sys_info.os == "Linux":
            try:
                # Get distribution information
                import distro
                sys_info["distribution"] = distro.name(pretty=True)
                sys_info["distribution_version"] = distro.version()
                sys_info["distribution_codename"] = distro.codename()
            except ImportError:
                # Fallback if distro module is not available
                try:
                    with open('/etc/os-release', 'r') as f:
                        os_release = dict(line.strip().split('=', 1) for line in f if '=' in line)
                    sys_info["distribution"] = os_release.get('PRETTY_NAME', '').strip('"')
                except:
                    pass
        return sys_info
    except Exception as e:
        logger.error(f"Error getting system info: {str(e)}")
        raise ValueError(f"Failed to get system information: {str(e)}")

def _format_size(size_bytes: int) -> str:
    """Helper function to format size in human-readable format"""
    if size_bytes == 0:
        return "0 B"
    size_names = ("B", "KB", "MB", "GB", "TB", "PB")
    i = int(math.floor(math.log(size_bytes, 1024)))
    p = math.pow(1024, i)
    s = round(size_bytes / p, 2)
    return f"{s} {size_names[i]}"
   
@register_tool
def file_status(
    directory: str = '.',
    pattern: str = '*'
) -> Dict[str, Any]:
    """
    Get status of files in a directory matching a pattern.

    Args:
        directory: Directory to scan (default: current directory)
        pattern: Glob pattern to match files (default: all files)
        
    Returns:
        Dictionary with file status information
    """
    logger.info(f"[tool calls] file_status directory: {directory}, pattern: {pattern}")
    try:
        # Resolve the path to handle relative paths and symlinks
        path = Path(directory).resolve()

        if not path.exists():
            raise ValueError(f"Directory does not exist: {directory}")
        
        if not path.is_dir():
            raise ValueError(f"Path is not a directory: {directory}")
        
        # get all files matching the pattern
        files = list(path.glob(pattern))

        # check if we should only count files in the current directory or recursively
        files = [file for file in files if file.is_file()]

        # calculate statistics
        total_size = sum(file.stat().st_size for file in files)

        # Count files by extension
        extensions = {}
        for file in files:
            ext = file.suffix.lower()
            if ext in extensions:
                extensions[ext] += 1
            else:
                extensions[ext] = 1

        # get the oldest and newest file
        if files:
            oldest_file = min(files, key=lambda f: f.stat().st_mtime)
            newest_file = max(files, key=lambda f: f.stat().st_mtime)
            oldest = {
                "path": str(oldest_file.relative_to(path.parent)),
                "modified": datetime.datetime.fromtimestamp(oldest_file.stat().st_mtime).isoformat()
            }
            newest = {
                "path": str(newest_file.relative_to(path.parent)),
                "modified": datetime.datetime.fromtimestamp(newest_file.stat().st_mtime).isoformat()
            }
        else:
            oldest = newest = None

        result = {
            'directory': str(path),
            'pattern': pattern,
            'total_size': total_size,
            'total_size_human': _format_size(total_size),
            'file_count': len(files),
            'extensions': extensions,
            'oldest_file': oldest,
            'newest_file': newest,
            'analyzed_time': datetime.datetime.now().isoformat()
        }
        return result

    except Exception as e:
        logger.error(f"Error analyzing files: {str(e)}")
        raise ValueError(f"Failed to analyze files in '{directory}': {str(e)}")

@register_tool
def web_search(query: str = None, num_results: int = 5) -> List[Dict[str, str]]:
    """
    Search the web for the given query and return a list of results.

    Args:
        query: The search query (default: None)
        num_results: Number of results to return (default: 5)
        
    Returns:
        List of dictionaries containing search results with 'title', 'link', and 'snippet'
    """
    logger.info(f"[tool calls] web_search query: {query}, num_results: {num_results}")

    try:
        # DuckDuckGo search API endpoint
        url = "https://serpapi.com/search"
        
        # Get API key from environment
        api_key = os.environ.get("SERPAPI_KEY")
        if not api_key:
            raise ValueError("SERPAPI_KEY environment variable not set")
        
        # Parameters for the search
        params = {
            'engine': 'duckduckgo',
            'q': query,
            'api_key': api_key,
            'kl': 'us-en'  # Region and language
        }
        # Send request
        response = requests.get(url, params=params)
        response.raise_for_status()
        
        # Parse the response
        data = response.json()
        
        # Extract organic results
        results = []
        if 'organic_results' in data:
            for result in data['organic_results'][:num_results]:
                results.append({
                    "title": result.get('title', ''),
                    "link": result.get('link', ''),
                    "snippet": result.get('snippet', '')
                })
        
        # Add knowledge graph if available
        if 'knowledge_graph' in data and len(results) < num_results:
            kg = data['knowledge_graph']
            results.append({
                "title": kg.get('title', ''),
                "link": kg.get('website', ''),
                "snippet": kg.get('description', '')
            })
        
        # Add related searches if needed
        if 'related_searches' in data and len(results) < num_results:
            for related in data['related_searches'][:num_results - len(results)]:
                results.append({
                    "title": f"Related: {related.get('query', '')}",
                    "link": related.get('link', ''),
                    "snippet": "Related search suggestion"
                })
        
        return results[:num_results]
    
    except Exception as e:
        logger.error(f"Search error: {str(e)}")
        raise ValueError(f"Failed to execute search '{query}': {str(e)}")
    
@register_tool
def http_request(
    url: str, 
    method: str = "GET", 
    headers: Optional[Dict[str, str]] = None, 
    data: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Make an HTTP request to the specified URL.

    Args:
        url: The URL to request
        method: HTTP method (default: "GET")
        headers: Optional dictionary of HTTP headers
        data: Optional dictionary of data to send in the request body
        
    Returns:
        Dictionary containing the response status, headers, and text
    """
    logger.info(f"[tool calls] http_request url: {url}, method: {method}, headers: {headers}, data: {data}")

    try:
        # Prepare request parameters
        kwargs = {
            "headers": headers or {}
        }
        
        # Handle different request methods with data
        if method.upper() in ["POST", "PUT", "PATCH"] and data:
            kwargs["json"] = data
            
        # Send request
        response = requests.request(method.upper(), url, **kwargs)
        
        # Prepare response data
        try:
            response_data = response.json()
        except:
            response_data = response.text
            
        return {
            "status_code": response.status_code,
            "headers": dict(response.headers),
            "data": response_data
        }
    
    except Exception as e:
        raise ValueError(f"Failed to execute HTTP request: {str(e)}")
    

@register_tool
def disk_usage(path: str = "/") -> Dict[str, Any]:
    """
    Get disk usage information for a specified path.
    
    Args:
        path: Path to check disk usage, default is root directory
        
    Returns:
        Dictionary containing disk usage information
    """
    logger.info(f"[tool calls] disk_usage path: {path}")
    try:
        usage = psutil.disk_usage(path)
        return {
            "path": path,
            "total_bytes": usage.total,
            "used_bytes": usage.used,
            "free_bytes": usage.free,
            "percent_used": usage.percent,
            "total_human": _format_size(usage.total),
            "used_human": _format_size(usage.used),
            "free_human": _format_size(usage.free),
            "updated_at": datetime.datetime.now().isoformat()
        }
    except Exception as e:
        logger.error(f"Error getting disk usage: {str(e)}")
        raise ValueError(f"Failed to get disk usage for '{path}': {str(e)}")
    
@register_tool
def process_list(limit: int = 10) -> List[Dict[str, Any]]:
    """
    Get list of running processes.
    
    Args:
        limit: Maximum number of processes to return, sorted by CPU usage
        
    Returns:
        List of process information dictionaries
    """
    logger.info(f"[tool calls] process_list limit: {limit}")
    try:
        processes = []
        for proc in psutil.process_iter(['pid', 'name', 'cpu_percent', 'memory_percent', 'username']):
            try:
                pinfo = proc.info
                processes.append({
                    "pid": pinfo['pid'],
                    "name": pinfo['name'],
                    "cpu_percent": pinfo['cpu_percent'],
                    "memory_percent": pinfo['memory_percent'],
                    "username": pinfo['username']
                })
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                continue
        
        # Sort by CPU usage and limit results
        processes.sort(key = lambda x: x['cpu_percent'], reverse = True)
        return processes[:limit]
    except Exception as e:
        logger.error(f"Error getting process list: {str(e)}")
        raise ValueError(f"Failed to get process list: {str(e)}")
    
@register_tool
def system_load() -> Dict[str, Any]:
    """
    Get system load informaton including: CPU, memory and disk usage

    Returns:
        Dictionary containing system load information
    """
    logger.info(f"[tool calls] system_load")
    try:
        # CPU usage
        cpu_percent = psutil.cpu_percent(interval=1)
        cpu_count = psutil.cpu_count()
        
        # Memory usage
        memory = psutil.virtual_memory()
        
        # Disk usage
        disk = psutil.disk_usage('/')
        
        return {
            "cpu": {
                "percent": cpu_percent,
                "count": cpu_count,
                "load_avg": psutil.getloadavg()
            },
            "memory": {
                "total": memory.total,
                "available": memory.available,
                "used": memory.used,
                "free": memory.free,
                "percent": memory.percent,
                "total_human": _format_size(memory.total),
                "available_human": _format_size(memory.available),
                "used_human": _format_size(memory.used),
                "free_human": _format_size(memory.free)
            },
            "disk": {
                "total": disk.total,
                "used": disk.used,
                "free": disk.free,
                "percent": disk.percent,
                "total_human": _format_size(disk.total),
                "used_human": _format_size(disk.used),
                "free_human": _format_size(disk.free)
            },
            "updated_at": datetime.datetime.now().isoformat()
        }
    except Exception as e:
        logger.error(f"Error getting system load: {str(e)}")
        raise ValueError(f"Failed to get system load information: {str(e)}")
    
@register_tool
def open_browser(url: str = None, q: str = None) -> str:
    """
    Open a URL in the default web browser.

    Args:
        url: The URL to open. Can also be a search query (will use Google).
        q: Search query (shorthand for url when it's a search term)

    Returns:
        Status message
    """
    # Determine the target from either url or q parameter
    target = url
    if not target and q:
        target = q

    if not target:
        raise ValueError("open_browser requires either 'url' or 'q' parameter")

    logger.info(f"[tool calls] open_browser target: {target}")

    # If it doesn't look like a URL, treat it as a search query
    if not target.startswith(('http://', 'https://', 'file://')):
        # 检查是否已经是 URL 编码过的，如果是，则解码后再编码
        import urllib.parse
        try:
            # 尝试解码，看看是否已经是编码过的
            decoded_target = urllib.parse.unquote(target)
            if decoded_target != target:
                target = decoded_target
        except:
            pass
        target = f"https://www.google.com/search?q={requests.utils.quote(target)}"

    try:
        webbrowser.open(target)
        return f"Opened browser with: {target}"
    except Exception as e:
        raise ValueError(f"Failed to open browser: {str(e)}")
    

@register_tool
def open_app(app_name: str) -> str:
    """
    Open an application by name.
    
    Args:
        app_name: Name of the application to open (e.g., 'notepad', 'code', 'chrome')
        
    Returns:
        Status message
    """
    logger.info(f"[tool calls] open_app app_name: {app_name}")
    system = platform.system()
    
    try:
        if system == "Windows":
            # Common Windows app mappings
            app_map = {
                'notepad': 'notepad.exe',
                'calculator': 'calc.exe',
                'explorer': 'explorer.exe',
                'cmd': 'cmd.exe',
                'powershell': 'powershell.exe',
                'chrome': 'chrome',
                'firefox': 'firefox',
                'edge': 'msedge',
                'code': 'code',
                'vscode': 'code',
            }
            cmd = app_map.get(app_name.lower(), app_name)
            subprocess.Popen(cmd, shell=True)
        elif system == "Darwin":  # macOS
            subprocess.Popen(['open', '-a', app_name])
        else:  # Linux
            subprocess.Popen([app_name], start_new_session=True)
        
        return f"Opened application: {app_name}"
    except Exception as e:
        raise ValueError(f"Failed to open application '{app_name}': {str(e)}")

@register_tool  
def clipboard_copy(text: str) -> str:
    """
    Copy text to system clipboard.
    
    Args:
        text: Text to copy to clipboard
        
    Returns:
        Status message
    """
    logger.info(f"[tool calls] clipboard_copy text: {text}")
    system = platform.system()
    
    try:
        if system == "Windows":
            subprocess.run(['clip'], input=text.encode('utf-16le'), check=True)
        elif system == "Darwin":  # macOS
            subprocess.run(['pbcopy'], input=text.encode('utf-8'), check=True)
        else:  # Linux
            # Try xclip first, then xsel
            try:
                subprocess.run(['xclip', '-selection', 'clipboard'], input=text.encode('utf-8'), check=True)
            except FileNotFoundError:
                subprocess.run(['xsel', '--clipboard', '--input'], input=text.encode('utf-8'), check=True)
        
        return f"Copied {len(text)} characters to clipboard"
    except Exception as e:
        raise ValueError(f"Failed to copy to clipboard: {str(e)}")

@register_tool
def create_docx(path: str, content: str, title: str = "") -> str:
    """
    Create a Word document (.docx) with the given content.
    Note: Requires python-docx package. If not installed, creates a .txt file instead.
    
    Args:
        path: Path for the output file (will add .docx extension if needed)
        content: Text content for the document
        title: Optional title for the document
        
    Returns:
        Status message with file path
    """
    # Ensure .docx extension
    if not path.lower().endswith('.docx'):
        path = path + '.docx'
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title if provided
        if title:
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Split content into paragraphs and add them
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                # Check if it's a heading (starts with #)
                if para_text.startswith('# '):
                    doc.add_heading(para_text[2:], level=1)
                elif para_text.startswith('## '):
                    doc.add_heading(para_text[3:], level=2)
                elif para_text.startswith('### '):
                    doc.add_heading(para_text[4:], level=3)
                else:
                    doc.add_paragraph(para_text.strip())
        
        doc.save(path)
        return f"Created Word document: {path}"
        
    except ImportError:
        # Fallback: create a text file with instructions
        txt_path = path.replace('.docx', '.txt')
        with open(txt_path, 'w', encoding='utf-8') as f:
            if title:
                f.write(f"{title}\n{'='*len(title)}\n\n")
            f.write(content)
        return f"python-docx not installed. Created text file instead: {txt_path}\nTo create .docx files, run: pip install python-docx"
